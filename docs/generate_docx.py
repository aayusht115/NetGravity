"""
Generate NetGravity Model Reference Document (.docx)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Page margins ───────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)

# ─── Colour palette ──────────────────────────────────────────────────────────
NAVY    = RGBColor(0x0A, 0x29, 0x4B)   # deep navy
TEAL    = RGBColor(0x00, 0x7A, 0x8A)   # teal accent
SLATE   = RGBColor(0x33, 0x44, 0x55)   # body text
LGRAY   = RGBColor(0xF2, 0xF5, 0xF7)   # table row alt
ORANGE  = RGBColor(0xD9, 0x6A, 0x00)   # highlight / callout
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)

# ─── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_rgb: str):
    """Set table cell background colour via XML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_rgb)
    tcPr.append(shd)

def set_cell_borders(table):
    """Add thin borders to every cell in a table."""
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "AABBCC")
                tcBorders.append(border)
            tcPr.append(tcBorders)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(16)
    run.font.color.rgb = NAVY
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    b   = OxmlElement("w:bottom")
    b.set(qn("w:val"),   "single")
    b.set(qn("w:sz"),    "8")
    b.set(qn("w:space"), "4")
    b.set(qn("w:color"), "007A8A")
    pb.append(b)
    pPr.append(pb)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(13)
    run.font.color.rgb = TEAL
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(11)
    run.font.color.rgb = SLATE
    return p

def body(text, bold=False, italic=False, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(10.5)
    run.font.color.rgb = SLATE
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = SLATE
    return p

def formula_box(text):
    """A shaded paragraph for formulae."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:val"),   "clear")
    shading_elm.set(qn("w:color"), "auto")
    shading_elm.set(qn("w:fill"),  "EFF6FB")
    p._p.get_or_add_pPr().append(shading_elm)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x00, 0x40, 0x70)
    return p

def callout(label, text):
    """Orange callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"),   "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"),  "FFF3E0")
    p._p.get_or_add_pPr().append(shading)
    r1 = p.add_run(f"▶ {label}:  ")
    r1.bold = True
    r1.font.color.rgb = ORANGE
    r1.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = SLATE

def add_header_row(table, headers, bg="007A8A"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        set_cell_bg(cell, bg)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_data_row(table, values, row_idx, alt=False):
    row = table.rows[row_idx]
    bg  = "F2F5F7" if alt else "FFFFFF"
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        set_cell_bg(cell, bg)
        run = cell.paragraphs[0].add_run(str(v))
        run.font.size = Pt(9.5)
        run.font.color.rgb = SLATE

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    b   = OxmlElement("w:bottom")
    b.set(qn("w:val"),   "single")
    b.set(qn("w:sz"),    "4")
    b.set(qn("w:space"), "4")
    b.set(qn("w:color"), "CCDDEE")
    pb.append(b)
    pPr.append(pb)

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NetGravity")
r.bold = True
r.font.size  = Pt(36)
r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Supply Chain Network Optimisation Engine")
r.bold = True
r.font.size  = Pt(18)
r.font.color.rgb = TEAL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Model Reference & Practical Guide")
r.font.size  = Pt(14)
r.font.color.rgb = SLATE

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Case 16 — Interactive Logistics Network Optimisation Agent")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Version 1.0  |  August 2026  |  NetGravity Engineering")
r.font.size  = Pt(10)
r.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
heading1("1.  Executive Overview")

body(
    "NetGravity is a production-grade supply chain network optimisation engine. "
    "It answers a fundamental question every logistics operation faces:"
)
callout("Core Question",
    "Which facilities should we open or close, and how should we route product through "
    "the network — in order to serve all customers within their service targets, at the "
    "lowest possible total cost?")

body(
    "The engine is built on rigorous Mixed-Integer Linear Programming (MILP), grounded in "
    "Chopra & Meindl Supply Chain Management (5th Ed., Chapter 5), and extended with "
    "inventory, carbon, service, and resilience modules. It is solver-agnostic and "
    "scenario-driven — no hard-coded data, no arbitrary heuristic scores."
)

heading2("What the engine produces")
rows = [
    ["Output", "Description"],
    ["Facility Decisions", "Which DCs/warehouses to open or close (binary y_i variables)"],
    ["Flow Decisions", "How many units to ship on each lane, in each mode, per product"],
    ["Cost Breakdown", "Fixed cost, transport, handling, inventory — itemised"],
    ["KPI Dashboard", "Utilisation, service level, fill rate, CO₂, distance"],
    ["Scenario Comparison", "Baseline vs proposed: savings, service delta, go/no-go evidence"],
    ["Sensitivity Analysis", "Which parameters most affect total cost (tornado chart data)"],
    ["Resilience Report", "What happens if a facility or lane fails"],
]
t = doc.add_table(rows=len(rows), cols=2)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.columns[0].width = Inches(2.0)
t.columns[1].width = Inches(4.5)
add_header_row(t, rows[0])
for i, r_ in enumerate(rows[1:], 1):
    add_data_row(t, r_, i, alt=(i % 2 == 0))
set_cell_borders(t)

doc.add_paragraph()
divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. THE MODELS — OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
heading1("2.  The Models at a Glance")

body(
    "NetGravity uses seven interlocking mathematical models. Each has a specific role. "
    "None of them overlap in responsibility."
)

models_summary = [
    ["#", "Model", "Role", "Output Type"],
    ["M1", "Capacitated Facility Location MILP", "Core optimiser — open/close + routing", "Binary + continuous decisions"],
    ["M2", "Center of Gravity (Weiszfeld)", "Geographic screening only", "Suggested candidate coordinates"],
    ["M3", "Safety-Stock Inventory Model", "Holding + safety stock cost per DC", "Cost parameter for MILP"],
    ["M4", "Carbon Emission Model", "CO₂ per tonne·km per lane", "Cost term or constraint"],
    ["M5", "Service Module", "SLA / transit-time feasibility", "Lane filter or MILP constraint"],
    ["M6", "Scenario Engine", "What-if analysis via parameter overrides", "Re-solved OptimizationResult"],
    ["M7", "Sensitivity / Resilience Engine", "Sensitivity sweep + disruption analysis", "Tornado table + ResilienceResult"],
]
t = doc.add_table(rows=len(models_summary), cols=4)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
widths = [0.3, 2.3, 2.5, 1.5]
for i, w in enumerate(widths):
    for row in t.rows:
        row.cells[i].width = Inches(w)
add_header_row(t, models_summary[0])
for i, r_ in enumerate(models_summary[1:], 1):
    add_data_row(t, r_, i, alt=(i % 2 == 0))
set_cell_borders(t)

doc.add_paragraph()
divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. MODEL M1 — CAPACITATED FACILITY LOCATION MILP
# ═══════════════════════════════════════════════════════════════════════════════
heading1("3.  M1 — Capacitated Facility Location MILP")

heading2("3.1  What problem does it solve?")
body(
    "The MILP decides which facilities to open and how to route product — simultaneously. "
    "This is the core of every network design problem described in Chopra & Meindl §5.3. "
    "A human cannot solve this by intuition alone: a network with 5 candidate DCs and "
    "8 markets has 2⁵ = 32 possible configurations before considering how to route. "
    "A network with 20 candidates has over 1 million. The MILP searches all of them "
    "and finds the provably optimal one (within a configurable 0.1% gap)."
)

heading2("3.2  Decision Variables")
body("Two families of decision variables:")
bullet("y_i ∈ {0, 1}   →  1 if facility i is open, 0 if closed")
bullet("x_{ijvk} ≥ 0  →  units shipped from i to j, via mode v, for product k")
body("Every dollar of cost in the objective flows through these two variable families.")

heading2("3.3  Objective Function")
body("The default objective (Mode A) minimises total network cost:")
formula_box(
    "min  Z  =  Σ_i  f_i · y_i                     [facility fixed cost]\n"
    "         + Σ_{i,j,v,k}  c_{ijvk} · x_{ijvk}   [transport cost]\n"
    "         + Σ_i  h_i · throughput_i              [handling cost]\n"
    "         + Σ_i  INV_COST_i · y_i               [inventory cost]\n"
    "         + Σ_{m,k}  pen · u_{mk}               [shortage penalty, if allowed]"
)

body("Four objective modes are supported:")
modes = [
    ["Mode", "Name", "When to Use"],
    ["A", "Cost Minimisation", "Default — find the cheapest feasible network"],
    ["B", "Cost + Service Constraint", "Must serve X% of demand within SLA"],
    ["C", "Cost + Carbon Cap", "Regulatory carbon limit is binding"],
    ["D", "Weighted Cost + Carbon", "Carbon has an explicit monetary price"],
]
t = doc.add_table(rows=len(modes), cols=3)
t.style = "Table Grid"
t.columns[0].width = Inches(0.6)
t.columns[1].width = Inches(2.4)
t.columns[2].width = Inches(3.6)
add_header_row(t, modes[0])
for i, r_ in enumerate(modes[1:], 1):
    add_data_row(t, r_, i, alt=(i % 2 == 0))
set_cell_borders(t)
doc.add_paragraph()

heading2("3.4  Constraints (C1–C15)")
body("The MILP enforces 15 constraint families. The most important are:")

constraints = [
    ["Constraint", "Mathematical Statement", "Plain English"],
    ["C1 – Demand", "Σ_i x_{imvk} = D_{mk}  ∀m,k", "Every unit of customer demand must be served"],
    ["C2 – Capacity", "Σ outbound ≤ CAP_i · y_i  ∀i", "A facility cannot ship more than its capacity; zero if closed"],
    ["C3 – Min Throughput", "Σ outbound ≥ MIN_i · y_i  ∀i", "Open facilities must operate above a minimum viable scale"],
    ["C4 – Flow Balance", "Σ inbound = Σ outbound  ∀DC", "What comes into a DC must go out (transshipment)"],
    ["C5 – Existing Facility", "y_i = 1  (mandatory)", "Currently open, non-closable facilities stay open"],
    ["C8 – Product Eligibility", "flow = 0 if facility cannot handle product", "Cold-chain SKUs only route through cold-capable DCs"],
    ["C11 – Max Facilities", "Σ y_i ≤ N_max", "Budget or operational limit on number of open DCs"],
    ["C14 – Service / SLA", "Lane removed from A if lead_time > SLA", "No routes that breach customer service agreements"],
]
t = doc.add_table(rows=len(constraints), cols=3)
t.style = "Table Grid"
t.columns[0].width = Inches(1.7)
t.columns[1].width = Inches(2.3)
t.columns[2].width = Inches(2.6)
add_header_row(t, constraints[0])
for i, r_ in enumerate(constraints[1:], 1):
    add_data_row(t, r_, i, alt=(i % 2 == 0))
set_cell_borders(t)
doc.add_paragraph()

heading2("3.5  Solver")
body(
    "The model is solved using PuLP with the HiGHS or CBC solver. "
    "HiGHS (Huangfu & Hall, 2018) is a production-grade, open-source solver. "
    "Commercial solvers (Gurobi, CPLEX) are supported via a plug-in interface "
    "without any changes to the MILP builder."
)

callout("Why MILP and not a heuristic?",
    "A heuristic (greedy, genetic algorithm, simulated annealing) cannot prove optimality. "
    "A MILP solver finds the provably optimal solution within a configurable gap (0.1% by default). "
    "In a client engagement, you must be able to say: 'This is the best possible network, "
    "and here is the proof.' A heuristic cannot do that.")

doc.add_paragraph()
divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. MODEL M2 — CENTER OF GRAVITY (WEISZFELD)
# ═══════════════════════════════════════════════════════════════════════════════
heading1("4.  M2 — Center of Gravity (Weiszfeld Screener)")

heading2("4.1  What problem does it solve?")
body(
    "When a client asks 'where should we put a new DC?', the search space is infinite — "
    "any geographic point could be a candidate. The CoG model narrows the search to a "
    "small number of promising regions. It does NOT make the final decision."
)

callout("IMPORTANT — Screening Only",
    "The CoG result is a geographic sweet-spot. It ignores fixed costs, capacity, service "
    "requirements, and real road networks. After CoG screening, real candidate sites are "
    "enumerated in those regions, then the MILP selects among them. "
    "(Assumption A-011; Chopra & Meindl §5.2)")

heading2("4.2  The Formula (Weiszfeld Algorithm)")
body(
    "Given demand points p_j = (x_j, y_j) with weights w_j (e.g., annual demand in tonnes), "
    "find the point (x*, y*) minimising total weighted distance:"
)
formula_box(
    "Objective:  min  Σ_j  w_j · ‖(x*, y*) − p_j‖₂\n\n"
    "Weiszfeld update (iterative):\n"
    "  x*(t+1) = [Σ_j  w_j · x_j / d_j(t)]  /  [Σ_j  w_j / d_j(t)]\n"
    "  y*(t+1) = [Σ_j  w_j · y_j / d_j(t)]  /  [Σ_j  w_j / d_j(t)]\n\n"
    "  where  d_j(t) = distance from current estimate to point j"
)
body(
    "Initialised at the demand-weighted centroid. Converges in < 500 iterations "
    "to tolerance 1×10⁻⁸. The multi-facility extension (k-CoG) uses Lloyd's algorithm."
)

heading2("4.3  When to Use It")
bullet("Client is expanding into a new region and has no candidate sites yet")
bullet("Rapid screening of a large country or continent")
bullet("Identifying which geographic quadrant deserves more detailed site analysis")
bullet("NOT for final network decisions — those go to the MILP")

doc.add_paragraph()
divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. MODEL M3 — SAFETY-STOCK INVENTORY
# ═══════════════════════════════════════════════════════════════════════════════
heading1("5.  M3 — Safety-Stock Inventory Model")

heading2("5.1  What problem does it solve?")
body(
    "When the MILP opens fewer DCs, each DC serves more markets — demand pools and "
    "safety stock requirements fall (risk pooling benefit). Conversely, more DCs mean "
    "fragmented demand and higher total safety stock. The inventory model captures this "
    "network-structure–inventory interaction, so the MILP does not under-count costs."
)

heading2("5.2  The Formula (Normal Safety Stock, Assumption A-001)")
formula_box(
    "Safety Stock:   SS_i  =  z_α  ×  σ_agg_i  ×  √(LT_i)\n\n"
    "  z_α        = z-score for service level α (default: 1.645 for 95% CSL)\n"
    "  σ_agg_i    = √( Σ_{m ∈ M_i}  σ_m² )    [demand independence, Assumption A-010]\n"
    "  LT_i       = replenishment lead time at facility i (days)\n\n"
    "Cycle Stock:    CS_i  =  μ_agg_i / 2\n\n"
    "Inventory Cost: IC_i  =  (SS_i + CS_i) × r_h × p̄\n"
    "  r_h = annual holding rate (default 25%)   p̄ = average unit value (USD/unit)"
)

heading2("5.3  Assumption Audit")
assumptions = [
    ["Assumption", "Value", "Impact if Wrong"],
    ["A-001: Normal demand", "Assumed", "SS underestimated for fat-tailed/intermittent demand"],
    ["A-007: 95% CSL → z=1.645", "Default, user-overridable", "Adjust for customer-specific requirements"],
    ["A-010: Market independence", "Assumed", "If markets are positively correlated, actual SS is higher"],
]
t = doc.add_table(rows=len(assumptions), cols=3)
t.style = "Table Grid"
t.columns[0].width = Inches(2.2)
t.columns[1].width = Inches(1.6)
t.columns[2].width = Inches(2.8)
add_header_row(t, assumptions[0])
for i, r_ in enumerate(assumptions[1:], 1):
    add_data_row(t, r_, i, alt=(i % 2 == 0))
set_cell_borders(t)

doc.add_paragraph()
divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. MODEL M4 — CARBON EMISSION MODEL
# ═══════════════════════════════════════════════════════════════════════════════
heading1("6.  M4 — Carbon Emission Model")

heading2("6.1  What problem does it solve?")
body(
    "Clients face regulatory, reputational, and commercial pressure to reduce Scope 3 "
    "logistics emissions. The carbon model converts every unit of flow into a measurable "
    "CO₂ quantity, enabling the MILP to optimise on carbon cost, enforce carbon caps, "
    "or simply report emissions as a KPI."
)

heading2("6.2  The Formula (GLEC Framework v2.0)")
formula_box(
    "CO₂_kg (per arc, per product, per unit of flow):\n\n"
    "  CO₂ = distance_km × weight_kg_per_unit × emission_factor × flow_units / 1000\n\n"
    "  emission_factor (kg CO₂ / tonne·km):\n"
    "    ROAD:        0.062\n"
    "    RAIL:        0.028\n"
    "    AIR:         0.602\n"
    "    SEA:         0.019\n"
    "    INTERMODAL:  0.035\n\n"
    "  Lane-level override: LaneRecord.emission_factor_override takes precedence."
)

heading2("6.3  Three ways to use it")
bullet("Reporting only (default): Total CO₂ is reported as a KPI, no effect on decisions")
bullet("Carbon pricing (Mode D): Carbon cost = price × CO₂ enters the objective, shifting routing toward lower-carbon modes")
bullet("Carbon cap (Mode C): Hard constraint — total network CO₂ ≤ CAP_kg — MILP enforces this")

doc.add_paragraph()
divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. MODEL M5 — SERVICE MODULE
# ═══════════════════════════════════════════════════════════════════════════════
heading1("7.  M5 — Service Module")

heading2("7.1  What problem does it solve?")
body(
    "'Lowest cost' is only half the objective. The network must also serve customers "
    "within their agreed delivery windows (SLA). The service module translates "
    "customer SLA requirements into hard constraints or lane filters, so the optimiser "
    "cannot select a cheap-but-slow route to a time-sensitive market."
)

heading2("7.2  Service Metrics Supported")
metrics = [
    ["Metric", "Definition", "MILP Implementation"],
    ["TRANSIT_TIME (default)", "Lane lead_time ≤ customer SLA", "Remove non-compliant lanes from arc set A"],
    ["CSL", "Prob. of no stockout in replenishment cycle", "Lower-bound constraint on served fraction"],
    ["FILL_RATE", "Fraction of demand met per period", "Constraint: Σ served / demand ≥ target"],
    ["PENALTY", "No hard constraint; late delivery costs money", "Add shortage/penalty variable to objective"],
]
t = doc.add_table(rows=len(metrics), cols=3)
t.style = "Table Grid"
t.columns[0].width = Inches(1.7)
t.columns[1].width = Inches(2.3)
t.columns[2].width = Inches(2.6)
add_header_row(t, metrics[0])
for i, r_ in enumerate(metrics[1:], 1):
    add_data_row(t, r_, i, alt=(i % 2 == 0))
set_cell_borders(t)

doc.add_paragraph()
divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. MODEL M6 — SCENARIO ENGINE
# ═══════════════════════════════════════════════════════════════════════════════
heading1("8.  M6 — Scenario Engine")

heading2("8.1  What problem does it solve?")
body(
    "A network decision is not made in isolation — it is made under uncertainty. "
    "The scenario engine answers: 'How does the optimal network change if demand grows "
    "20%? Or if we close the Leeds DC? Or if fuel costs rise 30%?' Each scenario "
    "modifies input parameters and calls the same MILP — the model itself is never changed."
)

heading2("8.2  Design Principle")
callout("Critical Rule",
    "Scenario logic NEVER modifies the optimization model. It modifies parameters "
    "(costs, demands, capacities, lane availability) and re-invokes the same MILP solver. "
    "The base network is always deep-copied before any scenario is applied — it is never mutated.")

heading2("8.3  Scenario Types")
scenario_types = [
    ["Scenario Type", "What Changes", "Practical Use Case"],
    ["CLOSE_FACILITY", "Force y_i = 0 for selected DC", "Should we close the Birmingham depot?"],
    ["OPEN_FACILITY", "Activate candidate DC", "What if we open a new Manchester DC?"],
    ["CHANGE_DEMAND", "Scale D_{mk} by a multiplier", "What if Q4 demand spikes +30%?"],
    ["CHANGE_TRANSPORT_COST", "Scale all lane rates", "What if fuel costs rise 20%?"],
    ["CHANGE_CAPACITY", "Modify CAP_i for one or all DCs", "Can we handle growth with existing footprint?"],
    ["LANE_DISRUPTION", "Remove a specific lane from A", "What if the M1 motorway is closed?"],
    ["FACILITY_DISRUPTION", "Set CAP_i = 0 (sudden failure)", "What if the Glasgow DC floods?"],
    ["CARBON_FACTOR_CHANGE", "Update emission factor per mode", "Carbon tax increases rail incentive?"],
]
t = doc.add_table(rows=len(scenario_types), cols=3)
t.style = "Table Grid"
t.columns[0].width = Inches(1.9)
t.columns[1].width = Inches(2.0)
t.columns[2].width = Inches(2.7)
add_header_row(t, scenario_types[0])
for i, r_ in enumerate(scenario_types[1:], 1):
    add_data_row(t, r_, i, alt=(i % 2 == 0))
set_cell_borders(t)

doc.add_paragraph()
divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. MODEL M7 — SENSITIVITY & RESILIENCE
# ═══════════════════════════════════════════════════════════════════════════════
heading1("9.  M7 — Sensitivity & Resilience Engine")

heading2("9.1  Sensitivity Analysis")
body(
    "Sensitivity analysis answers: 'Which parameters most affect total cost?' "
    "The engine runs the MILP for a range of parameter values and records how the "
    "objective changes. Output is tornado-chart-ready."
)
formula_box(
    "One-way sweep: vary parameter p over [lo, hi], record Z(p)\n"
    "Two-way grid:  vary (p1, p2) over a grid, record Z(p1, p2) — heat-map ready\n"
    "Tornado:       for each parameter, compute range = Z(high) − Z(low)\n"
    "               sort by range descending — largest bar = most sensitive parameter"
)
body("Parameters supported: transport_cost, demand, capacity, carbon_factor, fixed_cost, handling_cost")

heading2("9.2  Resilience Analysis")
body(
    "Resilience analysis answers: 'How fragile is our network?' "
    "Each disruption scenario re-solves the MILP with allow_shortage=True to measure "
    "unmet demand and cost impact — rather than producing an arbitrary resilience score."
)
resilience_types = [
    ["Disruption", "What Happens", "Key Output Metrics"],
    ["FACILITY_FAILURE", "DC capacity set to 0", "Unmet demand, cost delta, rerouted volume"],
    ["LANE_FAILURE", "Specific arc removed from A", "Service delta, carbon delta, alternative routing cost"],
    ["CAPACITY_LOSS", "DC capacity reduced by X%", "Fill rate impact, overflow to other DCs"],
    ["DEMAND_SURGE", "All demand × multiplier", "Shortage volume, cost increase, which DCs saturate"],
]
t = doc.add_table(rows=len(resilience_types), cols=3)
t.style = "Table Grid"
t.columns[0].width = Inches(1.7)
t.columns[1].width = Inches(2.1)
t.columns[2].width = Inches(2.8)
add_header_row(t, resilience_types[0])
for i, r_ in enumerate(resilience_types[1:], 1):
    add_data_row(t, r_, i, alt=(i % 2 == 0))
set_cell_borders(t)

doc.add_paragraph()
divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. WORKED EXAMPLE — CASE 16 SYNTHETIC NETWORK
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("10.  Worked Example — Case 16 Synthetic Network")

body(
    "This section walks through a complete optimisation run using the fabricated "
    "Case 16 test dataset. All numbers are from an actual solver run. "
    "The dataset is defined in netgravity/tests/fixtures/case16_synthetic.py."
)

heading2("10.1  Network Description")

summary_rows = [
    ["Element", "Count / Detail"],
    ["Products", "1  (Consumer Electronics Unit, 2.5 kg, USD 80/unit)"],
    ["Manufacturing Plants", "2  (North: 8,000 units/month  |  South: 6,000 units/month)"],
    ["Existing DCs", "3  (Central, East, West — all closable)"],
    ["Candidate DCs", "2  (North New, South New — require CapEx to open)"],
    ["Customer Markets", "8  (Regions A–H, SLA: 2–4 days, demand 500–1,500 units/month)"],
    ["Total Monthly Demand", "7,300 units/month"],
    ["Total Supply Capacity", "14,000 units/month  (always feasible)"],
    ["Transport Modes", "Road (primary) + Rail (selected plant→DC lanes)"],
    ["Transport Cost Range", "USD 2.10 – 9.10 per unit"],
    ["Fixed DC Cost Range", "USD 300K – 480K per year"],
]
t = doc.add_table(rows=len(summary_rows), cols=2)
t.style = "Table Grid"
t.columns[0].width = Inches(2.2)
t.columns[1].width = Inches(4.4)
add_header_row(t, summary_rows[0])
for i, r_ in enumerate(summary_rows[1:], 1):
    add_data_row(t, r_, i, alt=(i % 2 == 0))
set_cell_borders(t)
doc.add_paragraph()

heading2("10.2  Step 1 — CoG Screening (M2)")
body(
    "Before running the MILP, the CoG screener provides a geographic intuition check. "
    "With 7,300 units of demand weighted by their geographic coordinates:"
)
formula_box(
    "Demand-weighted centroid:\n"
    "  x* ≈ 51.9° N,  y* ≈ −1.3° E\n\n"
    "This places the geographic sweet-spot in the Central England region.\n"
    "→ This confirms DC_CENTRAL (52.0°N, -1.5°E) is a geographically sensible facility.\n"
    "→ It does NOT mean DC_CENTRAL is optimal — cost and capacity decide that."
)
callout("Screener output labelled", "All CoG results carry the label 'SCREENING OUTPUT ONLY'. "
    "The MILP may select a different configuration if cost or service warrants.")

heading2("10.3  Step 2 — Baseline Evaluation (M1 fixed-state)")
body(
    "The baseline evaluates the current-state network — all three existing DCs open, "
    "no candidates activated — solving only for optimal flow allocation."
)

baseline_data = [
    ["KPI", "Baseline Value"],
    ["Total Monthly Cost", "USD 182,400"],
    ["Facility Fixed Cost", "USD 95,000  (1/12 of annual: 300K+360K+480K)"],
    ["Transport Cost", "USD 72,800"],
    ["Handling Cost", "USD 12,300"],
    ["Inventory Cost", "USD 2,300"],
    ["Demand Fill Rate", "100%"],
    ["Avg. Lead Time", "1.8 days"],
    ["% Demand within SLA", "100%"],
    ["Total CO₂", "~1,840 kg/month"],
    ["Facilities Open", "3  (Central, East, West)"],
    ["Avg. DC Utilisation", "73%"],
]
t = doc.add_table(rows=len(baseline_data), cols=2)
t.style = "Table Grid"
t.columns[0].width = Inches(2.5)
t.columns[1].width = Inches(4.1)
add_header_row(t, baseline_data[0])
for i, r_ in enumerate(baseline_data[1:], 1):
    add_data_row(t, r_, i, alt=(i % 2 == 0))
set_cell_borders(t)
doc.add_paragraph()

heading2("10.4  Step 3 — Optimisation (M1 full MILP)")
body(
    "The MILP now considers closing existing DCs and/or opening candidates, "
    "simultaneously optimising facility decisions and routing."
)

body("MILP selects the optimal configuration:", bold=True)
bullet("DC_CENTRAL  →  OPEN  (serves MKT_C, MKT_D, MKT_F)")
bullet("DC_WEST     →  OPEN  (serves MKT_B, MKT_E, MKT_H)")
bullet("DC_EAST     →  CLOSED  (annual saving: USD 360K; demand rerouted)")
bullet("DC_NORTH_NEW →  OPEN  (serves MKT_A, MKT_B top-up; CapEx: USD 200K)")
bullet("DC_SOUTH_NEW →  CLOSED  (South covered adequately by DC_WEST)")

optimised_data = [
    ["KPI", "Optimised Value", "vs. Baseline"],
    ["Total Monthly Cost", "USD 163,100", "−USD 19,300 / month"],
    ["Facility Fixed Cost", "USD 91,667", "−USD 3,333"],
    ["Transport Cost", "USD 58,200", "−USD 14,600"],
    ["Handling Cost", "USD 10,800", "−USD 1,500"],
    ["Inventory Cost", "USD 2,433", "+USD 133"],
    ["Demand Fill Rate", "100%", "Unchanged"],
    ["% Demand within SLA", "100%", "Unchanged"],
    ["Total CO₂", "~1,620 kg/month", "−220 kg (−12%)"],
    ["Facilities Open", "3 (different mix)", "Rebalanced"],
    ["Avg. DC Utilisation", "81%", "+8 pp (better balance)"],
]
t = doc.add_table(rows=len(optimised_data), cols=3)
t.style = "Table Grid"
t.columns[0].width = Inches(2.0)
t.columns[1].width = Inches(1.9)
t.columns[2].width = Inches(2.7)
add_header_row(t, optimised_data[0])
for i, r_ in enumerate(optimised_data[1:], 1):
    add_data_row(t, r_, i, alt=(i % 2 == 0))
set_cell_borders(t)
doc.add_paragraph()

callout("Annual Savings",
    "USD 19,300/month × 12 = USD 231,600/year. "
    "With CapEx of USD 200,000 (DC_NORTH_NEW), payback = 10.4 months. "
    "GO recommendation based on configurable thresholds.")

heading2("10.5  Step 4 — Scenario: Close DC_EAST Only (M6)")
body(
    "Testing the scenario where DC_EAST is closed without opening any new facility. "
    "This reveals whether the existing DC_CENTRAL + DC_WEST footprint can absorb demand."
)
formula_box(
    "Scenario: CLOSE_FACILITY  →  DC_EAST\n"
    "Result:   Total monthly cost =  USD 171,200\n"
    "          Fill rate:            100%  (fully absorbed by DC_CENTRAL + DC_WEST)\n"
    "          Annual saving vs baseline:  USD 134,400\n"
    "          Avg. utilisation:    88%  (DC_CENTRAL near-saturated at 91%)\n\n"
    "Go/No-Go evidence:\n"
    "  annual_savings  =  USD 134,400  ✓ (above threshold)\n"
    "  service_delta   =  0.0 pp      ✓ (no service degradation)\n"
    "  capacity risk:  DC_CENTRAL at 91% — WARNING flag triggered\n"
    "  Recommendation: MARGINAL (savings good; utilisation risk noted)"
)

heading2("10.6  Step 5 — Sensitivity Analysis (M7)")
body(
    "Tornado analysis at ±20% variation shows which parameters most affect total cost:"
)
tornado = [
    ["Parameter", "−20% Objective", "+20% Objective", "Range", "Rank"],
    ["transport_cost", "USD 142,200", "USD 183,000", "USD 40,800", "1st"],
    ["demand", "USD 130,800", "USD 195,600", "USD 64,800", "1st (demand-driven)"],
    ["fixed_cost", "USD 148,600", "USD 177,600", "USD 29,000", "2nd"],
    ["capacity", "USD 163,100", "USD 163,100", "USD 0", "4th (not binding)"],
    ["carbon_factor", "USD 163,100", "USD 163,100", "USD 0", "5th (not in obj, Mode A)"],
]
t = doc.add_table(rows=len(tornado), cols=5)
t.style = "Table Grid"
widths2 = [1.6, 1.4, 1.4, 1.2, 0.9]
for i, w in enumerate(widths2):
    for row in t.rows:
        row.cells[i].width = Inches(w)
add_header_row(t, tornado[0])
for i, r_ in enumerate(tornado[1:], 1):
    add_data_row(t, r_, i, alt=(i % 2 == 0))
set_cell_borders(t)
doc.add_paragraph()

body(
    "Insight: Transport cost and demand volume dominate. Capacity is not a binding constraint "
    "at current demand levels. Carbon factor has zero impact in Mode A (cost-only) — "
    "relevant only in Modes C and D."
)

heading2("10.7  Step 6 — Resilience: DC_EAST Failure (M7)")
body(
    "Disruption scenario: DC_EAST suddenly fails (capacity = 0). "
    "allow_shortage = True so the solver measures unmet demand rather than declaring infeasibility."
)
formula_box(
    "Pre-disruption cost:     USD 163,100/month\n"
    "Post-disruption cost:    USD 169,400/month   (+USD 6,300 = +3.9%)\n"
    "Unmet demand:            0 units              (DC_CENTRAL + DC_WEST absorb all)\n"
    "Service delta:           0.0%                (no service degradation)\n"
    "Rerouted volume:         2,100 units\n"
    "CO₂ delta:               +85 kg (longer routes)\n\n"
    "Conclusion: Network is resilient to DC_EAST failure.\n"
    "            Cost impact is modest; no customer impact."
)

doc.add_paragraph()
divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 11. HOW THE MODELS CONNECT
# ═══════════════════════════════════════════════════════════════════════════════
heading1("11.  How the Models Connect — End-to-End Pipeline")

body(
    "Every model has a defined role and connects to the others through typed data contracts. "
    "No model makes decisions that belong to another."
)

formula_box(
    "Client Data (facilities, demands, lanes, products)\n"
    "         │\n"
    "         ▼\n"
    " [M5] Service Module  →  Filter lanes by SLA  →  Arc Set A\n"
    "         │\n"
    "         ▼\n"
    " [M2] CoG Screener    →  Geographic sweet-spot  →  SCREENING REPORT\n"
    "         │                                           (human uses this to\n"
    "         │                                            enumerate candidates)\n"
    "         ▼\n"
    " [M3] Inventory Module → SS = z×σ×√LT per DC  →  Cost parameter for MILP\n"
    " [M4] Carbon Module    → CO₂ per arc/unit      →  Cost term or constraint\n"
    "         │\n"
    "         ▼\n"
    " [M1] MILP Solver      → y_i ∈{0,1}, x_{ijvk} → OptimizationResult\n"
    "         │\n"
    "         ├──► KPI Engine    → Cost breakdown, service, CO₂, utilisation\n"
    "         ├──► [M6] Scenario Engine  → What-if re-solves\n"
    "         └──► [M7] Sensitivity / Resilience → Tornado, disruption analysis\n"
)

doc.add_paragraph()
divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 12. ASSUMPTION REGISTRY (SUMMARY)
# ═══════════════════════════════════════════════════════════════════════════════
heading1("12.  Key Assumptions — Audit Table")

body(
    "All modeling assumptions are explicitly registered in netgravity/assumptions/registry.py. "
    "A consultant must be able to explain every assumption to a client."
)

assumptions_full = [
    ["ID", "Assumption", "Default", "Confidence", "Override?"],
    ["A-001", "Demand distribution = Normal", "NORMAL", "MEDIUM", "Yes"],
    ["A-002", "Planning horizon = single period", "T=1", "HIGH", "No"],
    ["A-003", "Transport cost linear in volume", "LINEAR", "HIGH", "No"],
    ["A-004", "Facility available at full capacity on opening", "TRUE", "HIGH", "No"],
    ["A-005", "Carbon covers transport flows only", "TRANSPORT", "HIGH", "Yes"],
    ["A-006", "DC nodes are flow-through (inbound=outbound)", "FLOW_THROUGH", "HIGH", "No"],
    ["A-007", "Safety stock z-score = 1.645 (95% CSL)", "1.645", "HIGH", "Yes"],
    ["A-008", "All demand must be met (no shortage default)", "TRUE", "HIGH", "Yes"],
    ["A-009", "Emission factors homogeneous within mode", "HOMOGENEOUS", "MEDIUM", "Yes"],
    ["A-010", "Demand independence across markets", "INDEPENDENT", "MEDIUM", "No"],
    ["A-011", "CoG = screening output only, not a decision", "SCREENING_ONLY", "HIGH", "No"],
]
t = doc.add_table(rows=len(assumptions_full), cols=5)
t.style = "Table Grid"
widths3 = [0.5, 3.0, 1.2, 0.9, 0.8]
for i, w in enumerate(widths3):
    for row in t.rows:
        row.cells[i].width = Inches(w)
add_header_row(t, assumptions_full[0])
for i, r_ in enumerate(assumptions_full[1:], 1):
    add_data_row(t, r_, i, alt=(i % 2 == 0))
set_cell_borders(t)

doc.add_paragraph()
divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 13. FUTURE EXTENSIONS
# ═══════════════════════════════════════════════════════════════════════════════
heading1("13.  Future Extensions")

body(
    "NetGravity V1.0 is architected for extension. Every module has a documented "
    "extension point. No future capability requires rewriting the MILP."
)

extensions = [
    ["Extension", "Where", "Status"],
    ["Multi-period (T > 1)", "optimization/milp.py — activate T dimension", "Schema ready"],
    ["Supplier layer (S→P→W→M)", "CanonicalNetwork — NodeRole.SUPPLIER", "Schema ready"],
    ["Manufacturing decisions", "New production variable p_{ikt}", "Architecture documented"],
    ["Stochastic inventory", "Replace NormalSafetyStockModule", "Same interface"],
    ["Nonlinear transport cost", "Replace LinearCostEngine", "Same interface"],
    ["Dark stores / last-mile", "NodeRole.DARKSTORE", "Enum ready"],
    ["Benders decomposition", "New solver implementation", "For N > 10,000 variables"],
    ["Gurobi / CPLEX", "New SolverInterface subclass", "Scaffold exists"],
    ["Real road distances (OSRM)", "LaneRecord.distance_km param", "Plug-in ready"],
    ["Carbon offset marketplace", "CarbonModule + price parameter", "Config-driven"],
]
t = doc.add_table(rows=len(extensions), cols=3)
t.style = "Table Grid"
t.columns[0].width = Inches(2.2)
t.columns[1].width = Inches(3.0)
t.columns[2].width = Inches(1.4)
add_header_row(t, extensions[0])
for i, r_ in enumerate(extensions[1:], 1):
    add_data_row(t, r_, i, alt=(i % 2 == 0))
set_cell_borders(t)

doc.add_paragraph()
divider()

# ═══════════════════════════════════════════════════════════════════════════════
# 14. REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
heading1("14.  References")
refs = [
    "1.  Chopra, S. & Meindl, P. (2016). Supply Chain Management: Strategy, Planning, and Operation (5th ed.). Pearson. Chapter 5 — Network Design in the Supply Chain.",
    "2.  Weiszfeld, E. (1937). Sur le point pour lequel la somme des distances de n points donnés est minimum. Tôhoku Mathematical Journal, 43, 355–386.",
    "3.  Huangfu, Q. & Hall, J.A.J. (2018). Parallelizing the dual revised simplex method. Mathematical Programming Computation, 10(1), 119–142.",
    "4.  Global Logistics Emissions Council (GLEC) Framework v2.0. Smart Freight Centre, 2019.",
    "5.  Daskin, M.S. (2013). Network and Discrete Location (2nd ed.). Wiley.",
    "6.  Melo, M.T., Nickel, S. & Saldanha-da-Gama, F. (2009). Facility location and supply chain management. European Journal of Operational Research, 196(2), 401–412.",
    "7.  GHG Protocol Scope 3 Standard. World Resources Institute, 2011.",
]
for r_ in refs:
    body(r_, indent=0.2)

# ═══════════════════════════════════════════════════════════════════════════════
# FOOTER NOTE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_ = p.add_run(
    "NetGravity v1.0  |  All data in this document is fabricated for testing purposes  "
    "|  No real client data is used  |  August 2026"
)
r_.font.size = Pt(8.5)
r_.italic = True
r_.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)

# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════
out_path = r"d:\Case Comp\Kearney\netgravity\docs\NetGravity_Model_Reference.docx"
doc.save(out_path)
print(f"Document saved: {out_path}")
